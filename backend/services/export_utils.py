"""Export utilities for PDF, DOCX, XLSX, CSV, and JSON formats."""

import csv
import io
import json
from typing import Any


def to_json(data: Any) -> bytes:
    return json.dumps(data, indent=2, default=str).encode("utf-8")


def to_csv(rows: list[dict], columns: list[str] | None = None) -> bytes:
    if not rows:
        return b""
    cols = columns or list(rows[0].keys())
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=cols, extrasaction="ignore")
    writer.writeheader()
    writer.writerows(rows)
    return buf.getvalue().encode("utf-8")


def to_xlsx(rows: list[dict], sheet_name: str = "Export") -> bytes:
    try:
        from openpyxl import Workbook
    except ImportError:
        return to_csv(rows)

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31]
    if not rows:
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    cols = list(rows[0].keys())
    ws.append(cols)
    for row in rows:
        ws.append([row.get(c) for c in cols])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def to_docx(title: str, content: str) -> bytes:
    try:
        from docx import Document
    except ImportError:
        return f"{title}\n\n{content}".encode("utf-8")

    doc = Document()
    doc.add_heading(title, 0)
    for para in content.split("\n"):
        if para.strip():
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(title: str, content: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return f"{title}\n\n{content}".encode("utf-8")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, title[:80])
    y -= 30
    c.setFont("Helvetica", 10)
    for line in content.split("\n"):
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 50
        c.drawString(50, y, line[:100])
        y -= 14
    c.save()
    return buf.getvalue()
